"""
檔案處理與內容擷取服務
支援 PDF、DOCX、TXT、Markdown 檔案的載入與分塊處理
"""
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path

# PDF 處理
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# DOCX 處理
try:
    from docx import Document
except ImportError:
    Document = None

from app.settings import settings


logger = logging.getLogger(__name__)


class FileIngestService:
    """檔案內容擷取服務類別"""
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.md', '.markdown'}
        self.chunk_size = 1000  # 預設分塊大小
        self.chunk_overlap = 200  # 分塊重疊大小
    
    def is_supported_file(self, file_path: Union[str, Path]) -> bool:
        """檢查檔案是否為支援的格式"""
        path = Path(file_path)
        return path.suffix.lower() in self.supported_extensions
    
    async def load_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        載入檔案內容
        
        Args:
            file_path: 檔案路徑
            
        Returns:
            包含檔案內容的字典
        """
        path = Path(file_path)
        
        # 檢查檔案是否存在
        if not path.exists():
            # 嘗試在上傳目錄中尋找
            upload_path = self.upload_dir / path.name
            if upload_path.exists():
                path = upload_path
            else:
                return {
                    "ok": False,
                    "reason": "file_not_found",
                    "message": f"檔案不存在: {file_path}",
                    "data": None
                }
        
        # 檢查檔案格式
        if not self.is_supported_file(path):
            return {
                "ok": False,
                "reason": "unsupported_format",
                "message": f"不支援的檔案格式: {path.suffix}",
                "data": None
            }
        
        try:
            # 根據檔案類型載入內容
            if path.suffix.lower() == '.pdf':
                content = await self._load_pdf(path)
            elif path.suffix.lower() == '.docx':
                content = await self._load_docx(path)
            elif path.suffix.lower() in {'.txt', '.md', '.markdown'}:
                content = await self._load_text(path)
            else:
                raise ValueError(f"未實作的檔案類型: {path.suffix}")
            
            return {
                "ok": True,
                "data": {
                    "file_path": str(path),
                    "file_name": path.name,
                    "file_size": path.stat().st_size,
                    "file_type": path.suffix.lower(),
                    "content": content["text"],
                    "metadata": content.get("metadata", {}),
                    "page_count": content.get("page_count", 1)
                },
                "source": "file_ingest",
                "timestamp": None
            }
            
        except Exception as e:
            logger.error(f"檔案載入失敗: {str(e)} - {path}")
            return {
                "ok": False,
                "reason": "load_failed",
                "message": f"檔案載入失敗: {str(e)}",
                "data": None
            }
    
    async def _load_pdf(self, path: Path) -> Dict[str, Any]:
        """載入 PDF 檔案內容"""
        if not PdfReader:
            raise ImportError("pypdf 套件未安裝，無法處理 PDF 檔案")
        
        reader = PdfReader(str(path))
        text_content = []
        metadata = {
            "pages": []
        }
        
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text.strip():
                text_content.append(f"[頁面 {page_num}]\n{page_text}")
                metadata["pages"].append({
                    "page_number": page_num,
                    "text_length": len(page_text),
                    "has_content": bool(page_text.strip())
                })
        
        # 取得 PDF 元資料
        if reader.metadata:
            metadata.update({
                "title": reader.metadata.get("/Title", ""),
                "author": reader.metadata.get("/Author", ""),
                "subject": reader.metadata.get("/Subject", ""),
                "creator": reader.metadata.get("/Creator", "")
            })
        
        return {
            "text": "\n\n".join(text_content),
            "metadata": metadata,
            "page_count": len(reader.pages)
        }
    
    async def _load_docx(self, path: Path) -> Dict[str, Any]:
        """載入 DOCX 檔案內容"""
        if not Document:
            raise ImportError("python-docx 套件未安裝，無法處理 DOCX 檔案")
        
        doc = Document(str(path))
        text_content = []
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        }
        
        # 擷取段落內容
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # 擷取表格內容
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text.append(" | ".join(row_text))
            if table_text:
                text_content.append("\n".join(table_text))
        
        # 取得文件屬性
        if hasattr(doc, 'core_properties'):
            props = doc.core_properties
            metadata.update({
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else ""
            })
        
        return {
            "text": "\n\n".join(text_content),
            "metadata": metadata,
            "page_count": 1  # DOCX 沒有頁面概念
        }
    
    async def _load_text(self, path: Path) -> Dict[str, Any]:
        """載入純文字檔案內容"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 計算基本統計
        lines = content.split('\n')
        metadata = {
            "lines": len(lines),
            "characters": len(content),
            "words": len(content.split()),
            "encoding": "utf-8"
        }
        
        return {
            "text": content,
            "metadata": metadata,
            "page_count": 1
        }
    
    def chunk_text(self, 
                   text: str, 
                   chunk_size: Optional[int] = None,
                   chunk_overlap: Optional[int] = None,
                   preserve_structure: bool = True) -> List[Dict[str, Any]]:
        """
        將文字分塊處理
        
        Args:
            text: 要分塊的文字
            chunk_size: 分塊大小（字元數）
            chunk_overlap: 分塊重疊大小
            preserve_structure: 是否保持結構（段落、句子邊界）
            
        Returns:
            分塊結果列表
        """
        if chunk_size is None:
            chunk_size = self.chunk_size
        if chunk_overlap is None:
            chunk_overlap = self.chunk_overlap
        
        chunks = []
        
        if preserve_structure:
            # 按段落分割
            paragraphs = text.split('\n\n')
            current_chunk = ""
            chunk_id = 0
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                # 如果當前段落加入後超過大小限制
                if len(current_chunk) + len(para) > chunk_size and current_chunk:
                    chunks.append({
                        "id": chunk_id,
                        "text": current_chunk.strip(),
                        "start_char": len("".join([c["text"] for c in chunks])),
                        "end_char": len("".join([c["text"] for c in chunks])) + len(current_chunk),
                        "length": len(current_chunk.strip())
                    })
                    chunk_id += 1
                    
                    # 保留重疊部分
                    if chunk_overlap > 0:
                        overlap_text = current_chunk[-chunk_overlap:]
                        current_chunk = overlap_text + "\n\n" + para
                    else:
                        current_chunk = para
                else:
                    if current_chunk:
                        current_chunk += "\n\n" + para
                    else:
                        current_chunk = para
            
            # 處理最後一個分塊
            if current_chunk.strip():
                chunks.append({
                    "id": chunk_id,
                    "text": current_chunk.strip(),
                    "start_char": len("".join([c["text"] for c in chunks])),
                    "end_char": len("".join([c["text"] for c in chunks])) + len(current_chunk),
                    "length": len(current_chunk.strip())
                })
        
        else:
            # 簡單的固定大小分塊
            for i in range(0, len(text), chunk_size - chunk_overlap):
                chunk_text = text[i:i + chunk_size]
                if chunk_text.strip():
                    chunks.append({
                        "id": len(chunks),
                        "text": chunk_text.strip(),
                        "start_char": i,
                        "end_char": min(i + chunk_size, len(text)),
                        "length": len(chunk_text.strip())
                    })
        
        return chunks
    
    async def process_file(self, 
                          file_path: Union[str, Path],
                          chunk_size: Optional[int] = None,
                          chunk_overlap: Optional[int] = None) -> Dict[str, Any]:
        """
        完整處理檔案：載入 + 分塊
        
        Args:
            file_path: 檔案路徑
            chunk_size: 分塊大小
            chunk_overlap: 分塊重疊大小
            
        Returns:
            處理結果字典
        """
        # 載入檔案
        load_result = await self.load_file(file_path)
        if not load_result["ok"]:
            return load_result
        
        file_data = load_result["data"]
        
        # 分塊處理
        chunks = self.chunk_text(
            file_data["content"],
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        
        return {
            "ok": True,
            "data": {
                "file_info": {
                    "path": file_data["file_path"],
                    "name": file_data["file_name"],
                    "size": file_data["file_size"],
                    "type": file_data["file_type"],
                    "page_count": file_data["page_count"]
                },
                "content": file_data["content"],
                "metadata": file_data["metadata"],
                "chunks": chunks,
                "chunk_count": len(chunks),
                "total_length": len(file_data["content"])
            },
            "source": "file_ingest",
            "timestamp": None
        }


# 全域檔案處理服務實例
file_ingest_service = FileIngestService()
